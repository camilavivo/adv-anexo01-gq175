import re, unicodedata
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _norm(s: str) -> str:
    s = unicodedata.normalize('NFKC', s or '')
    s = s.replace('\xa0', ' ')
    s = re.sub(r'\s+', ' ', s).strip().lower()
    s = s.replace('→','->').replace(' / ','/')
    return s

def _format_cell(cell):
    text = cell.text
    cell.text = ''
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rf = run.font
    rf.name = 'Arial'
    rf.size = Pt(9)
    rf.color.rgb = RGBColor(0,0,0)

def _list_to_map(pairs: List[dict]) -> Dict[str,str]:
    out = {}
    for p in pairs:
        out[_norm(p['item'])] = p['descricao']
    return out

def fill_docx(template_path: str, t3_pairs: List[dict], t4_pairs: List[dict]) -> Tuple[Document, dict]:
    doc = Document(template_path)
    t3 = _list_to_map(t3_pairs)
    t4 = _list_to_map(t4_pairs)
    report = {'tabela3': {'preenchidos': [], 'nao_encontrados': []}, 'tabela4': {'preenchidos': [], 'nao_encontrados': []}}

    def try_fill(table, mapa, bucket):
        max_cols = max(len(r.cells) for r in table.rows)
        if max_cols < 3:
            return 0
        hits = 0
        seen = set()
        for r in table.rows:
            if len(r.cells) < 3: continue
            label_raw = r.cells[1].text
            key = _norm(label_raw)
            if key in mapa:
                r.cells[2].text = mapa[key]
                _format_cell(r.cells[2])
                report[bucket]['preenchidos'].append(label_raw.strip())
                hits += 1
                seen.add(key)
        for k in mapa.keys():
            if k not in seen:
                report[bucket]['nao_encontrados'].append(k)
        return hits
    t3h = t4h = 0
    for table in doc.tables:
        if t3h == 0:
            t3h = try_fill(table, t3, 'tabela3')
        if t4h == 0:
            t4h = try_fill(table, t4, 'tabela4')
        if t3h and t4h:
            break
    return doc, report
